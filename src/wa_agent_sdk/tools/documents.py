"""Document ingestion: PDF, DOCX, Markdown, and plain-text formats."""

from __future__ import annotations

import html.parser
import io
import mimetypes
import os
from dataclasses import dataclass
from typing import Callable

from ..exceptions import UnsupportedDocumentError


@dataclass(slots=True)
class ParsedDocument:
    filename: str
    mime: str
    text: str
    pages: int | None = None
    truncated: bool = False


class _HTMLTextExtractor(html.parser.HTMLParser):
    _SKIP = {"script", "style", "head", "title"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._chunks: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in self._SKIP:
            self._skip_depth += 1
        if tag in ("p", "br", "div", "li", "tr", "h1", "h2", "h3"):
            self._chunks.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP and self._skip_depth:
            self._skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if not self._skip_depth and data.strip():
            self._chunks.append(data)

    def text(self) -> str:
        raw = "".join(self._chunks)
        lines = [ln.strip() for ln in raw.splitlines()]
        return "\n".join(ln for ln in lines if ln)


def _parse_pdf(data: bytes) -> tuple[str, int | None]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise UnsupportedDocumentError("PDF support requires 'pip install pypdf'") from exc
    reader = PdfReader(io.BytesIO(data))
    page_texts: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            content = (page.extract_text() or "").strip()
        except Exception:  # noqa: BLE001 - single broken page must not kill the doc
            content = ""
        if content:
            page_texts.append(f"[page {index}]\n{content}")
    joined = "\n\n".join(page_texts)
    total_pages = len(reader.pages)
    if not joined.strip():
        joined = (
            f"(This {total_pages}-page PDF has no extractable text layer — "
            "it is probably scanned images. Tell the user you cannot read its contents.)"
        )
    return joined, total_pages


def _parse_docx(data: bytes) -> tuple[str, int | None]:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise UnsupportedDocumentError("DOCX support requires 'pip install python-docx'") from exc
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        parts.append("\n".join(rows))
    return "\n\n".join(parts), None


def _parse_plain(data: bytes) -> tuple[str, int | None]:
    return data.decode("utf-8", errors="replace"), None


def _parse_html(data: bytes) -> tuple[str, int | None]:
    extractor = _HTMLTextExtractor()
    extractor.feed(data.decode("utf-8", errors="replace"))
    return extractor.text(), None


_PARSERS: dict[str, Callable[[bytes], tuple[str, int | None]]] = {
    "pdf": _parse_pdf,
    "docx": _parse_docx,
    "txt": _parse_plain,
    "md": _parse_plain,
    "markdown": _parse_plain,
    "rst": _parse_plain,
    "csv": _parse_plain,
    "tsv": _parse_plain,
    "json": _parse_plain,
    "yaml": _parse_plain,
    "yml": _parse_plain,
    "xml": _parse_plain,
    "log": _parse_plain,
    "ini": _parse_plain,
    "py": _parse_plain,
    "js": _parse_plain,
    "ts": _parse_plain,
    "html": _parse_html,
    "htm": _parse_html,
}

SUPPORTED_EXTENSIONS = sorted(_PARSERS)


def parse_document(
    filename: str,
    data: bytes,
    *,
    max_chars: int = 15_000,
) -> ParsedDocument:
    """Parse *data* according to the extension of *filename*.

    Raises :class:`UnsupportedDocumentError` for unknown or unparseable types.
    """
    ext = os.path.splitext(filename)[1].lower().lstrip(".")
    parser = _PARSERS.get(ext)
    if parser is None:
        raise UnsupportedDocumentError(
            f"Unsupported document type '.{ext}'. Supported: {', .'.join(SUPPORTED_EXTENSIONS)}"
        )
    try:
        text, pages = parser(data)
    except UnsupportedDocumentError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise UnsupportedDocumentError(f"Failed to parse '{filename}': {exc}") from exc

    mime = mimetypes.guess_type(filename)[0] or "application/octet-stream"
    truncated = False
    if len(text) > max_chars:
        cut_note = f"\n\n[... document truncated, showing first {max_chars} of {len(text)} characters]"
        text = text[:max_chars] + cut_note
        truncated = True
    return ParsedDocument(filename=filename, mime=mime, text=text, pages=pages, truncated=truncated)


def format_document_block(doc: ParsedDocument) -> str:
    meta = [f'filename="{doc.filename}"', f'mime="{doc.mime}"']
    if doc.pages:
        meta.append(f"pages={doc.pages}")
    header = ", ".join(meta)
    return f"<document {header}>\n{doc.text}\n</document>"
