"""Core SDK tests: config, tools, documents, memory, agent construction."""

import asyncio
import base64
import io

import pytest

from wa_agent_sdk import (
    LLMConfig,
    ToolRegistry,
    WhatsAppAgent,
    create_provider,
    parse_document,
    prepare_image,
    tool,
)
from wa_agent_sdk.exceptions import UnsupportedDocumentError, WaAgentError
from wa_agent_sdk.memory import ConversationMemory
from wa_agent_sdk.llm.base import ChatMessage
from wa_agent_sdk.tools.builtin import calculate


def test_provider_resolution_and_typo_suggestions():
    cfg = LLMConfig(provider="groq", model="llama-3.3-70b-versatile", api_key="t")
    assert cfg.kind == "openai" and "groq.com" in cfg.resolved_base_url
    with pytest.raises(WaAgentError, match="Did you mean: openai"):
        LLMConfig(provider="opnai", model="x").info
    assert type(create_provider(cfg)).__name__ == "OpenAICompatibleProvider"


@pytest.mark.parametrize(
    ("provider", "base", "env"),
    [
        ("openai", "https://api.openai.com/v1", "OPENAI_API_KEY"),
        ("groq", "https://api.groq.com/openai/v1", "GROQ_API_KEY"),
        ("deepseek", "https://api.deepseek.com/v1", "DEEPSEEK_API_KEY"),
        ("openrouter", "https://openrouter.ai/api/v1", "OPENROUTER_API_KEY"),
        ("nvidia", "https://integrate.api.nvidia.com/v1", "NVIDIA_API_KEY"),
        ("xai", "https://api.x.ai/v1", "XAI_API_KEY"),
        ("grok", "https://api.x.ai/v1", "XAI_API_KEY"),
        ("anthropic", "https://api.anthropic.com", "ANTHROPIC_API_KEY"),
        ("gemini", "https://generativelanguage.googleapis.com", "GEMINI_API_KEY"),
        ("ollama", "http://localhost:11434/v1", ""),
    ],
)
def test_known_provider_table(provider, base, env):
    from wa_agent_sdk.config import _KNOWN_PROVIDERS

    info = _KNOWN_PROVIDERS[provider]
    assert info["base_url"] == base and info["env"] == env


def test_calculator():
    assert calculate("(2+3)*4") == "(2+3)*4 = 20"
    assert calculate("sqrt(16) + 2**10").endswith("= 1028")
    with pytest.raises(WaAgentError):
        calculate("__import__('os')")


def test_tool_schema_generation_and_execution():
    @tool(description="Greet someone.")
    def greet(name: str, times: int = 1) -> str:
        return "hi " * times

    reg = ToolRegistry()
    reg.register(greet)
    schema = reg.schemas()[0]
    assert schema["parameters"]["properties"]["times"]["type"] == "integer"
    assert schema["parameters"]["required"] == ["name"]
    assert asyncio.run(reg.get("greet").run({"name": "sam", "times": 2})) == "hi hi "


def test_markdown_parsing():
    md = parse_document("notes.md", b"# Title\n\nHello world")
    assert "Title" in md.text and not md.truncated


def test_docx_parsing_or_clean_env_error():
    try:
        import docx  # noqa: F401

        d = docx.Document()
        d.add_paragraph("Paragraph one")
        t = d.add_table(rows=1, cols=2)
        t.rows[0].cells[0].text = "A"
        t.rows[0].cells[1].text = "B"
        buf = io.BytesIO()
        d.save(buf)
        parsed = parse_document("report.docx", buf.getvalue())
        assert "Paragraph one" in parsed.text and "A | B" in parsed.text
    except ImportError:
        # lxml blocked on some machines (e.g. Application Control policy)
        with pytest.raises(UnsupportedDocumentError):
            parse_document("report.docx", b"PK\x03\x04 broken")


def test_pdf_roundtrip_and_broken_pdf():
    from pypdf import PdfWriter

    w = PdfWriter()
    w.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    w.write(buf)
    pdf = parse_document("blank.pdf", buf.getvalue())
    assert pdf.pages == 1 and "no extractable text layer" in pdf.text
    with pytest.raises(UnsupportedDocumentError):
        parse_document("broken.pdf", b"%PDF-1.4 fake")


def test_image_preparation_downscales():
    from PIL import Image

    img = Image.new("RGB", (3000, 1000), (200, 30, 30))
    b = io.BytesIO()
    img.save(b, "PNG")
    b64, mime = prepare_image(b.getvalue(), max_side=800)
    size = Image.open(io.BytesIO(base64.b64decode(b64))).size
    assert max(size) <= 800 and mime == "image/jpeg"


def test_memory_trimming():
    mem = ConversationMemory(max_messages=4, max_chars=5000)
    for i in range(10):
        mem.append("jid", ChatMessage(role="user", content=f"m{i}"))
    h = mem.history("jid")
    assert len(h) == 4 and h[-1].content == "m9"


def test_agent_construction_with_builtin_tools():
    agent = WhatsAppAgent(llm=LLMConfig(provider="gemini", model="gemini-2.0-flash", api_key="k"))
    for expected in ("calculate", "current_datetime", "remember_note"):
        assert expected in agent.tools.names()
    sp = agent._system_prompt()
    assert "Current date/time" in sp and "calculate" in sp


def test_notes_roundtrip_scoped_by_chat():
    from wa_agent_sdk import context

    agent = WhatsAppAgent(llm=LLMConfig(provider="gemini", model="g", api_key="k"))

    async def run():
        token = context.current_chat_jid.set("123@s.whatsapp.net")
        r1 = await agent.tools.get("remember_note").run({"note": "user likes tea"})
        r2 = await agent.tools.get("recall_notes").run({})
        context.current_chat_jid.reset(token)
        return r1, r2

    r1, r2 = asyncio.run(run())
    assert "Noted" in r1 and "tea" in r2


def test_dict_style_construction():
    agent = WhatsAppAgent(llm={"provider": "ollama", "model": "llama3"}, system_prompt="x")
    assert agent.config.llm.model == "llama3"
